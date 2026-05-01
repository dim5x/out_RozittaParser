"""
features/export/generator.py — Генерация DOCX / JSON / MD / HTML из архива Telegram

Читает данные из DBManager и создаёт файлы экспорта.
Все XML-трюки (закладки, ссылки) делегированы в features/export/xml_magic.py.

Поддерживаемые режимы (split_mode) для DocxGenerator:
    "none"  — один файл со всеми сообщениями
    "day"   — по одному файлу на каждый день
    "month" — по одному файлу на каждый месяц
    "post"  — по одному файлу на каждый пост (+ комментарии если include_comments)

Нет Qt-зависимостей. Нет Telethon-зависимостей.
Весь код — синхронный (python-docx не поддерживает async).

Пример использования (в ExportWorker.run()):
    with DBManager(cfg.db_path) as db:
        gen = DocxGenerator(db, output_dir=params.output_dir)
        files = gen.generate(
            chat_id          = result.chat_id,
            chat_title       = result.chat_title,
            split_mode       = params.split_mode,
            topic_id         = params.topic_id,
            user_id          = params.user_id,
            include_comments = params.download_comments,
            period_label     = result.period_label,
            log              = self.log_message.emit,
        )

ПРИМЕЧАНИЕ по merge_group_id:
    DBManager.get_messages() должен возвращать колонки merge_group_id (индекс 15)
    и merge_part_index (индекс 16). Если SELECT их не включает — добавьте в database.py:
        SELECT ..., merge_group_id, merge_part_index FROM messages ...
    До этого генераторы работают корректно: merge-поля читаются защищённо через len(row).
"""

from __future__ import annotations

import html as html_lib
import json
import logging
import os
import re
from collections import defaultdict
from pathlib import Path
from typing import Callable, Dict, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from config import VALID_SPLIT_MODES
from core.database import DBManager
from core.exceptions import DocxGenerationError, EmptyDataError
from core.utils import sanitize_filename, is_image_path
from features.export import xml_magic

logger = logging.getLogger(__name__)

_LogCallback = Callable[[str], None]

# Ширина изображений в документе (дюймы)
_IMAGE_WIDTH_INCHES = 4.5
# Отступ комментариев (дюймы)
_COMMENT_INDENT_INCHES = 0.3
# Линия-разделитель между сообщениями
_SEPARATOR = "─" * 60


# ==============================================================================
# Индексы колонок из DBManager.get_messages() / get_post_with_comments()
# ==============================================================================
# SELECT id, chat_id, message_id, topic_id, user_id, username,
#        date, text, media_path, file_type, file_size,
#        reply_to_msg_id, post_id, is_comment, linked_chat_id,
#        merge_group_id, merge_part_index
# FROM messages
#
_COL_ID             = 0
_COL_CHAT_ID        = 1
_COL_MESSAGE_ID     = 2
_COL_TOPIC_ID       = 3
_COL_USER_ID        = 4
_COL_USERNAME       = 5
_COL_DATE           = 6
_COL_TEXT           = 7
_COL_MEDIA_PATH     = 8
_COL_FILE_TYPE      = 9
_COL_FILE_SIZE      = 10
_COL_REPLY_TO       = 11
_COL_POST_ID        = 12
_COL_IS_COMMENT     = 13
_COL_LINKED_CHAT    = 14
_COL_MERGE_GROUP_ID = 15   # ← задача 5: merge groups
_COL_MERGE_PART_IDX = 16   # ← задача 5: порядок части внутри группы


# ==============================================================================
# Вспомогательные функции (модульный уровень)
# ==============================================================================

def _word_count(text: Optional[str]) -> int:
    """Быстрый подсчёт слов (split по пробелам)."""
    return len(text.split()) if text else 0


def _get_merge_group_id(row: tuple) -> Optional[int]:
    """
    Защищённое чтение merge_group_id.
    Возвращает None если колонка отсутствует в SELECT.
    """
    return row[_COL_MERGE_GROUP_ID] if len(row) > _COL_MERGE_GROUP_ID else None


def _group_by_merge(messages: List[tuple]) -> List[List[tuple]]:
    """
    Группирует последовательные сообщения с одинаковым merge_group_id
    в единые блоки для отображения.

    Одиночные сообщения (merge_group_id=None) образуют группы из одного элемента.
    Сообщения с одинаковым merge_group_id подряд объединяются в одну группу.

    Args:
        messages: Строки из DBManager.get_messages().

    Returns:
        Список групп. Каждая группа — список из 1+ строк.
    """
    groups: List[List[tuple]] = []
    i = 0
    while i < len(messages):
        merge_id = _get_merge_group_id(messages[i])
        if merge_id is not None:
            # Собираем все подряд идущие части с тем же merge_group_id
            group = [messages[i]]
            j = i + 1
            while j < len(messages) and _get_merge_group_id(messages[j]) == merge_id:
                group.append(messages[j])
                j += 1
            groups.append(group)
            i = j
        else:
            groups.append([messages[i]])
            i += 1
    return groups


def _topic_suffix(topic_id: Optional[int]) -> str:
    """Возвращает строку '_topicN' или '' если topic_id is None."""
    return f"_topic{topic_id}" if topic_id is not None else ""


# ==============================================================================
# DocxGenerator
# ==============================================================================

class DocxGenerator:
    """
    Генератор DOCX-документов из архива Telegram.

    Читает строки из DBManager (сохранённые через ParserService),
    форматирует их в Word-документы со структурой, закладками и ссылками.

    Args:
        db:         Открытый DBManager (контекстный менеджер управляет им снаружи).
        output_dir: Корневая папка для сохранения DOCX-файлов.
    """

    def __init__(self, db: DBManager, output_dir: str = "output") -> None:
        self._db          = db
        self._output_dir  = output_dir
        self._log: _LogCallback = logger.info

        # Текущий контекст (устанавливается в generate() для _build_path)
        self._chat_title:   str           = "chat"
        self._period_label: str           = "fullchat"
        self._topic_id:     Optional[int] = None   # ← задача 1
        # Транскрипции: {message_id: text} — загружаются в generate()
        self._transcriptions: Dict[int, str] = {}

    # ------------------------------------------------------------------
    # Публичный API
    # ------------------------------------------------------------------

    def generate(
        self,
        chat_id:          int,
        chat_title:       str           = "",
        split_mode:       str           = "none",
        topic_id:         Optional[int] = None,
        user_id:          Optional[int] = None,
        include_comments: bool          = False,
        period_label:     str           = "fullchat",
        date_from:        Optional[str] = None,
        date_to:          Optional[str] = None,
        log:              _LogCallback  = None,
    ) -> List[str]:
        """
        Главный метод: генерирует DOCX в соответствии с split_mode.

        Returns:
            Список путей к созданным DOCX-файлам. Пустой список при ошибке.

        Raises:
            DocxGenerationError: при критической ошибке создания файла.
            EmptyDataError:      если в БД нет сообщений для данного chat_id.
        """
        if split_mode not in VALID_SPLIT_MODES:
            raise DocxGenerationError(
                file_path="",
                message=f"Неверный split_mode: '{split_mode}'. "
                        f"Допустимые: {VALID_SPLIT_MODES}",
            )

        self._log          = log or logger.info
        self._chat_title   = chat_title or f"chat_{chat_id}"
        self._period_label = period_label
        self._topic_id     = topic_id   # ← задача 1: сохраняем для _build_path
        os.makedirs(self._output_dir, exist_ok=True)

        # Загружаем все транскрипции чата одним запросом
        try:
            self._transcriptions = self._db.get_transcriptions_for_chat(chat_id)
            if self._transcriptions:
                self._log(f"🎙 STT: загружено {len(self._transcriptions)} транскрипций")
        except Exception:
            self._transcriptions = {}

        self._log(f"📄 Генерация DOCX (режим: {split_mode})...")
        logger.info(
            "export: generate chat_id=%s split=%s topic=%s user=%s comments=%s",
            chat_id, split_mode, topic_id, user_id, include_comments,
        )

        try:
            if split_mode == "post":
                files = self._generate_by_posts(
                    chat_id          = chat_id,
                    topic_id         = topic_id,
                    user_id          = user_id,
                    include_comments = include_comments,
                    date_from        = date_from,
                    date_to          = date_to,
                )
            else:
                messages = self._db.get_messages(
                    chat_id          = chat_id,
                    topic_id         = topic_id,
                    user_id          = user_id,
                    include_comments = include_comments,
                    date_from        = date_from,
                    date_to          = date_to,
                )
                if not messages:
                    raise EmptyDataError(chat_id, topic_id)

                if split_mode == "day":
                    files = self._generate_by_day(messages, chat_id)
                elif split_mode == "month":
                    files = self._generate_by_month(messages, chat_id)
                else:  # "none"
                    files = self._generate_single(messages, chat_id)

        except (EmptyDataError, DocxGenerationError):
            raise
        except Exception as exc:
            logger.exception("export: unexpected error in generate()")
            raise DocxGenerationError(
                file_path=self._output_dir,
                message=f"Неожиданная ошибка генерации: {exc}",
                original=exc,
            ) from exc

        self._log(f"✅ Создано файлов: {len(files)}")
        return files

    # ------------------------------------------------------------------
    # Режимы генерации
    # ------------------------------------------------------------------

    def _generate_single(
        self,
        messages: List[tuple],
        chat_id:  int,
    ) -> List[str]:
        """Режим "none" — один файл со всеми сообщениями."""
        xml_magic.reset_counter()
        doc = Document()

        title = doc.add_heading(f"Архив чата: {self._chat_title}", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for group in _group_by_merge(messages):       # ← задача 5
            self._add_group_to_doc(doc, group)

        file_path = self._build_path("archive")
        self._save_doc(doc, file_path)
        self._log(f"  ✅ {os.path.basename(file_path)} ({len(messages)} сообщений)")
        return [file_path]

    def _generate_by_day(
        self,
        messages: List[tuple],
        chat_id:  int,
    ) -> List[str]:
        """Режим "day" — по одному файлу на каждый день."""
        days: Dict[str, List[tuple]] = defaultdict(list)
        for msg in messages:
            day = msg[_COL_DATE][:10]   # "YYYY-MM-DD"
            days[day].append(msg)

        files: List[str] = []
        for day, day_msgs in sorted(days.items()):
            xml_magic.reset_counter()
            doc = Document()

            title = doc.add_heading(f"Сообщения за {day}", level=1)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            for group in _group_by_merge(day_msgs):   # ← задача 5
                self._add_group_to_doc(doc, group)

            file_path = self._build_path(f"day_{day}")
            self._save_doc(doc, file_path)
            files.append(file_path)
            self._log(f"  ✅ {os.path.basename(file_path)} ({len(day_msgs)} сообщений)")

        self._log(f"📅 Создано {len(files)} файлов по дням")
        return files

    def _generate_by_month(
        self,
        messages: List[tuple],
        chat_id:  int,
    ) -> List[str]:
        """Режим "month" — по одному файлу на каждый месяц."""
        months: Dict[str, List[tuple]] = defaultdict(list)
        for msg in messages:
            month = msg[_COL_DATE][:7]  # "YYYY-MM"
            months[month].append(msg)

        files: List[str] = []
        for month, month_msgs in sorted(months.items()):
            xml_magic.reset_counter()
            doc = Document()

            title = doc.add_heading(f"Сообщения за {month}", level=1)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            for group in _group_by_merge(month_msgs):  # ← задача 5
                self._add_group_to_doc(doc, group)

            file_path = self._build_path(f"month_{month}")
            self._save_doc(doc, file_path)
            files.append(file_path)
            self._log(f"  ✅ {os.path.basename(file_path)} ({len(month_msgs)} сообщений)")

        self._log(f"📅 Создано {len(files)} файлов по месяцам")
        return files

    def _generate_by_posts(
        self,
        chat_id:          int,
        topic_id:         Optional[int],
        user_id:          Optional[int],
        include_comments: bool,
        date_from:        Optional[str] = None,
        date_to:          Optional[str] = None,
    ) -> List[str]:
        """Режим "post" — по одному файлу на каждый пост."""
        posts = self._db.get_messages(
            chat_id          = chat_id,
            topic_id         = topic_id,
            user_id          = user_id,
            include_comments = False,
            date_from        = date_from,
            date_to          = date_to,
        )
        if not posts:
            raise EmptyDataError(chat_id, topic_id)

        self._log(f"📝 Найдено постов: {len(posts)}")
        files: List[str] = []

        for post in posts:
            post_id  = post[_COL_MESSAGE_ID]
            xml_magic.reset_counter()
            doc = Document()

            title = doc.add_heading(f"Пост #{post_id}", level=1)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Сам пост
            self._add_message_to_doc(doc, post, is_post=True)

            # Комментарии (если нужны)
            comments: list = []
            if include_comments:
                all_rows = self._db.get_post_with_comments(chat_id, post_id)
                comments = [
                    r for r in all_rows
                    if not (r[_COL_MESSAGE_ID] == post_id and r[_COL_IS_COMMENT] == 0)
                ]
                if comments:
                    doc.add_paragraph()
                    comment_header = doc.add_heading(
                        f"💬 Комментарии ({len(comments)})", level=2
                    )
                    comment_header.paragraph_format.space_before = Pt(12)

                    for group in _group_by_merge(comments):   # ← задача 5
                        self._add_group_to_doc(doc, group, is_comment=True)

            file_path = self._build_path(f"post_{post_id}")
            self._save_doc(doc, file_path)
            files.append(file_path)

            comment_count = len(comments)
            suffix = f" ({comment_count} комментариев)" if include_comments else ""
            self._log(f"  ✅ {os.path.basename(file_path)}{suffix}")

        return files

    # ------------------------------------------------------------------
    # Группы сообщений (merge)  — задача 5
    # ------------------------------------------------------------------

    def _add_group_to_doc(
        self,
        doc:        Document,
        group:      List[tuple],
        is_post:    bool = False,
        is_comment: bool = False,
    ) -> None:
        """
        Добавляет группу сообщений в документ.

        Если группа содержит одно сообщение — вызывает _add_message_to_doc.
        Если несколько (merge group) — выводит единый заголовок и склеенный текст,
        без повторяющихся разделителей между частями.
        """
        if len(group) == 1:
            self._add_message_to_doc(doc, group[0], is_post=is_post, is_comment=is_comment)
            return

        first      = group[0]
        msg_id     = first[_COL_MESSAGE_ID]
        username   = first[_COL_USERNAME] or "Unknown"
        date_str   = first[_COL_DATE] or ""
        reply_to   = first[_COL_REPLY_TO]

        # --- Закладка (якорь по первому сообщению группы) ---
        anchor_p = doc.add_paragraph()
        xml_magic.add_bookmark(anchor_p, f"msg_{msg_id}")
        anchor_p.paragraph_format.space_before = Pt(0)
        anchor_p.paragraph_format.space_after  = Pt(0)

        # --- Единый заголовок ---
        header_p = doc.add_paragraph()
        if is_post:
            header_text = f"📌 ПОСТ от {username}"
            font_size   = Pt(12)
            font_color  = RGBColor(0, 102, 204)
        elif is_comment:
            header_text = f"  💬 {username}"
            font_size   = Pt(10)
            font_color  = RGBColor(102, 102, 102)
        else:
            header_text = f"👤 {username}"
            font_size   = Pt(11)
            font_color  = RGBColor(51, 51, 51)

        run = header_p.add_run(header_text)
        run.bold           = True
        run.font.size      = font_size
        run.font.color.rgb = font_color

        date_run = header_p.add_run(f"\n📅 {date_str}")
        date_run.font.size      = Pt(9)
        date_run.font.color.rgb = RGBColor(128, 128, 128)

        if is_comment:
            header_p.paragraph_format.left_indent = Inches(_COMMENT_INDENT_INCHES)

        # --- Ссылка-ответ (только у первой части) ---
        if reply_to:
            reply_p = doc.add_paragraph()
            reply_p.add_run("↩️ В ответ на: ")
            xml_magic.add_internal_hyperlink(reply_p, reply_to, f"сообщение #{reply_to}")
            if is_comment:
                reply_p.paragraph_format.left_indent = Inches(_COMMENT_INDENT_INCHES)

        # --- Текст и медиа каждой части (без заголовков и разделителей) ---
        for part in group:
            text       = (part[_COL_TEXT] or "").strip()
            media_path = part[_COL_MEDIA_PATH]
            part_id    = part[_COL_MESSAGE_ID]

            if text:
                text_p = doc.add_paragraph()
                xml_magic.write_text_with_links(text_p, text)
                if is_comment:
                    text_p.paragraph_format.left_indent = Inches(_COMMENT_INDENT_INCHES)

            if media_path and os.path.exists(media_path):
                abs_path = os.path.abspath(media_path)
                media_p  = doc.add_paragraph("📎 Медиафайл: ")
                file_uri = Path(abs_path).as_uri()
                xml_magic.add_external_hyperlink(
                    media_p, file_uri, os.path.basename(abs_path)
                )
                if is_comment:
                    media_p.paragraph_format.left_indent = Inches(_COMMENT_INDENT_INCHES)

                if is_image_path(abs_path):
                    try:
                        img_p   = doc.add_paragraph()
                        img_run = img_p.add_run()
                        img_run.add_picture(abs_path, width=Inches(_IMAGE_WIDTH_INCHES))
                        if is_comment:
                            img_p.paragraph_format.left_indent = Inches(_COMMENT_INDENT_INCHES)
                    except Exception as exc:
                        logger.warning(
                            "export: cannot insert image %s: %s",
                            os.path.basename(abs_path), exc,
                        )

            elif media_path:
                media_p = doc.add_paragraph(
                    f"📎 [медиафайл недоступен]: {os.path.basename(media_path)}"
                )
                if is_comment:
                    media_p.paragraph_format.left_indent = Inches(_COMMENT_INDENT_INCHES)

            # STT для этой части
            file_type = part[_COL_FILE_TYPE] or ""
            if file_type in ("voice", "video_note", "videomessage") and part_id in self._transcriptions:
                stt_text = self._transcriptions[part_id]
                if stt_text:
                    stt_p     = doc.add_paragraph()
                    stt_label = stt_p.add_run("🎙 Распознанная речь: ")
                    stt_label.bold           = True
                    stt_label.font.size      = Pt(10)
                    stt_label.font.color.rgb = RGBColor(80, 80, 80)
                    stt_run                  = stt_p.add_run(stt_text)
                    stt_run.font.size        = Pt(10)
                    stt_run.font.color.rgb   = RGBColor(40, 40, 40)
                    stt_run.italic           = True
                    if is_comment:
                        stt_p.paragraph_format.left_indent = Inches(_COMMENT_INDENT_INCHES)

        # --- Единый разделитель в конце всей группы ---
        if not is_comment:
            sep_p = doc.add_paragraph(_SEPARATOR)
            sep_p.paragraph_format.space_before = Pt(4)
            sep_p.paragraph_format.space_after  = Pt(4)
            for run in sep_p.runs:
                run.font.color.rgb = RGBColor(200, 200, 200)

    # ------------------------------------------------------------------
    # Добавление одного сообщения в документ
    # ------------------------------------------------------------------

    def _add_message_to_doc(
        self,
        doc:        Document,
        msg:        tuple,
        is_post:    bool = False,
        is_comment: bool = False,
    ) -> None:
        """
        Форматирует одну строку из БД и добавляет её в документ.

        Структура блока сообщения:
            [Параграф с закладкой]
            [Заголовок: имя + дата]
            [↩️ В ответ на: ссылка] (если reply_to не None)
            [Текст сообщения]        (с автоссылками)
            [📎 Медиафайл: ссылка]   (если есть медиа)
            [Изображение]            (если медиа — картинка)
            [──────────────]         (разделитель, не для комментариев)
        """
        msg_id     = msg[_COL_MESSAGE_ID]
        username   = msg[_COL_USERNAME]  or "Unknown"
        date_str   = msg[_COL_DATE]      or ""
        text       = msg[_COL_TEXT]      or ""
        media_path = msg[_COL_MEDIA_PATH]
        reply_to   = msg[_COL_REPLY_TO]

        # --- Закладка ---
        anchor_p = doc.add_paragraph()
        xml_magic.add_bookmark(anchor_p, f"msg_{msg_id}")
        anchor_p.paragraph_format.space_before = Pt(0)
        anchor_p.paragraph_format.space_after  = Pt(0)

        # --- Заголовок сообщения ---
        header_p = doc.add_paragraph()
        if is_post:
            header_text = f"📌 ПОСТ от {username}"
            font_size   = Pt(12)
            font_color  = RGBColor(0, 102, 204)
        elif is_comment:
            header_text = f"  💬 {username}"
            font_size   = Pt(10)
            font_color  = RGBColor(102, 102, 102)
        else:
            header_text = f"👤 {username}"
            font_size   = Pt(11)
            font_color  = RGBColor(51, 51, 51)

        run = header_p.add_run(header_text)
        run.bold             = True
        run.font.size        = font_size
        run.font.color.rgb   = font_color

        date_run = header_p.add_run(f"\n📅 {date_str}")
        date_run.font.size      = Pt(9)
        date_run.font.color.rgb = RGBColor(128, 128, 128)

        if is_comment:
            header_p.paragraph_format.left_indent = Inches(_COMMENT_INDENT_INCHES)

        # --- Ссылка-ответ ---
        if reply_to:
            reply_p = doc.add_paragraph()
            reply_p.add_run("↩️ В ответ на: ")
            xml_magic.add_internal_hyperlink(
                reply_p, reply_to, f"сообщение #{reply_to}"
            )
            if is_comment:
                reply_p.paragraph_format.left_indent = Inches(_COMMENT_INDENT_INCHES)

        # --- Текст с авто-ссылками ---
        if text:
            text_p = doc.add_paragraph()
            xml_magic.write_text_with_links(text_p, text)
            if is_comment:
                text_p.paragraph_format.left_indent = Inches(_COMMENT_INDENT_INCHES)

        # --- Медиа ---
        if media_path and os.path.exists(media_path):
            abs_path = os.path.abspath(media_path)
            media_p  = doc.add_paragraph("📎 Медиафайл: ")
            file_uri = Path(abs_path).as_uri()
            xml_magic.add_external_hyperlink(
                media_p, file_uri, os.path.basename(abs_path)
            )
            if is_comment:
                media_p.paragraph_format.left_indent = Inches(_COMMENT_INDENT_INCHES)

            if is_image_path(abs_path):
                try:
                    img_p   = doc.add_paragraph()
                    img_run = img_p.add_run()
                    img_run.add_picture(abs_path, width=Inches(_IMAGE_WIDTH_INCHES))
                    if is_comment:
                        img_p.paragraph_format.left_indent = Inches(_COMMENT_INDENT_INCHES)
                except Exception as exc:
                    logger.warning(
                        "export: cannot insert image %s: %s",
                        os.path.basename(abs_path), exc,
                    )
                    self._log(
                        f"⚠️ Не удалось вставить изображение "
                        f"{os.path.basename(abs_path)}: {exc}"
                    )

        elif media_path:
            media_p = doc.add_paragraph(
                f"📎 [медиафайл недоступен]: {os.path.basename(media_path)}"
            )
            if is_comment:
                media_p.paragraph_format.left_indent = Inches(_COMMENT_INDENT_INCHES)

        # --- Транскрипция ---
        file_type = msg[_COL_FILE_TYPE] or ""
        if file_type in ("voice", "video_note", 'videomessage') and msg_id in self._transcriptions:
            stt_text = self._transcriptions[msg_id]
            if stt_text:
                stt_p     = doc.add_paragraph()
                stt_label = stt_p.add_run("🎙 Распознанная речь: ")
                stt_label.bold           = True
                stt_label.font.size      = Pt(10)
                stt_label.font.color.rgb = RGBColor(80, 80, 80)
                stt_run                  = stt_p.add_run(stt_text)
                stt_run.font.size        = Pt(10)
                stt_run.font.color.rgb   = RGBColor(40, 40, 40)
                stt_run.italic           = True
                if is_comment:
                    stt_p.paragraph_format.left_indent = Inches(_COMMENT_INDENT_INCHES)

        # --- Разделитель (не для комментариев) ---
        if not is_comment:
            sep_p = doc.add_paragraph(_SEPARATOR)
            sep_p.paragraph_format.space_before = Pt(4)
            sep_p.paragraph_format.space_after  = Pt(4)
            for run in sep_p.runs:
                run.font.color.rgb = RGBColor(200, 200, 200)

    # ------------------------------------------------------------------
    # Вспомогательные методы
    # ------------------------------------------------------------------

    def _build_path(self, kind: str) -> str:
        """
        Строит полный путь к DOCX-файлу.

        Формат имени (задача 1 — topic_id добавлен):
            <chat>_topic{N}_<kind>_<period>.docx   (если topic_id задан)
            <chat>_<kind>_<period>.docx             (если topic_id=None)

        Args:
            kind: Тип файла ("archive", "day_2025-01-15", "post_42" и т.д.).

        Returns:
            Абсолютный путь к файлу.
        """
        safe_title = sanitize_filename(self._chat_title)
        topic_sfx  = _topic_suffix(self._topic_id)
        filename   = f"{safe_title}{topic_sfx}_{kind}_{self._period_label}.docx"
        return os.path.join(self._output_dir, filename)

    def _save_doc(self, doc: Document, file_path: str) -> None:
        """
        Сохраняет документ на диск.

        Raises:
            DocxGenerationError: если не удалось записать файл.
        """
        try:
            doc.save(file_path)
            logger.debug("export: saved %s", file_path)
        except Exception as exc:
            logger.error("export: save failed %s: %s", file_path, exc)
            raise DocxGenerationError(
                file_path = file_path,
                message   = f"Не удалось сохранить файл: {exc}",
                original  = exc,
            ) from exc


# ==============================================================================
# JsonGenerator
# ==============================================================================

class JsonGenerator:
    """
    Генерирует JSON-архив переписки из SQLite.

    Структура выходного файла — плоский список объектов, пригодный для
    загрузки в NotebookLM / ChatGPT и простого чтения вручную.

    При ai_split=True файл разбивается на части по ai_split_chunk_words слов.

    Нет Qt-зависимостей. Нет Telethon-зависимостей. Только stdlib + DBManager.
    """

    def __init__(self, db: DBManager, output_dir: str = "output") -> None:
        self._db         = db
        self._output_dir = output_dir

    def generate(
        self,
        chat_id:              int,
        chat_title:           str,
        *,
        topic_id:             Optional[int]  = None,      # ← задача 2
        user_id:              Optional[int]  = None,
        include_comments:     bool           = False,
        ai_split:             bool           = False,
        period_label:         str            = "fullchat", # ← задача 3
        ai_split_chunk_words: int            = 300_000,    # ← задача 4
        date_from:            Optional[str]  = None,
        date_to:              Optional[str]  = None,
        log:                  _LogCallback   = lambda _: None,
    ) -> List[str]:
        """
        Основная точка входа. Строит JSON и сохраняет на диск.

        Имя файла:
            <chat>_topic{N}_{period}_history.json   (с topic_id)
            <chat>_{period}_history.json            (без topic_id)
            <chat>_topic{N}_{period}_history_part_{k}.json   (с ai_split)

        Returns:
            Список абсолютных путей к созданным .json файлам.

        Raises:
            EmptyDataError: нет сообщений для экспорта.
            OSError: ошибка записи файла.
        """
        os.makedirs(self._output_dir, exist_ok=True)

        log("📋 Загружаю сообщения из БД для JSON-экспорта...")
        rows = self._db.get_messages(
            chat_id,
            topic_id         = topic_id,
            user_id          = user_id,
            include_comments = include_comments,
            date_from        = date_from,
            date_to          = date_to,
        )

        if not rows:
            raise EmptyDataError(f"Нет сообщений для чата {chat_id}")

        log(f"📊 Строк получено: {len(rows)}")

        stt_map:   dict[int, str] = self._db.get_transcriptions_for_chat(chat_id)
        safe_title = sanitize_filename(chat_title)
        topic_sfx  = _topic_suffix(topic_id)           # ← задача 2
        base_name  = f"{safe_title}{topic_sfx}_{period_label}_history"  # ← задача 3

        # ── Без разбивки: один файл ────────────────────────────────────
        if not ai_split:
            records: List[dict] = []
            for row in rows:
                msg_id = row[_COL_MESSAGE_ID]
                records.append(self._make_record(row, stt_map.get(msg_id)))
            out_path = os.path.join(self._output_dir, f"{base_name}.json")
            self._write_json(out_path, records, log)
            return [out_path]

        # ── С разбивкой по ai_split_chunk_words слов ───────────────────
        out_paths: List[str] = []
        chunk:      List[dict] = []
        words      = 0
        part       = 1

        for row in rows:
            msg_id  = row[_COL_MESSAGE_ID]
            record  = self._make_record(row, stt_map.get(msg_id))
            chunk.append(record)
            words += _word_count(record.get("text")) + _word_count(record.get("stt_text"))

            if words >= ai_split_chunk_words:               # ← задача 4
                path = os.path.join(
                    self._output_dir,
                    f"{base_name}_part_{part}.json",
                )
                self._write_json(path, chunk, log)
                out_paths.append(path)
                chunk = []
                words = 0
                part += 1

        if chunk:   # последний (возможно неполный) чанк
            path = os.path.join(
                self._output_dir,
                f"{base_name}_part_{part}.json",
            )
            self._write_json(path, chunk, log)
            out_paths.append(path)

        log(f"✅ JSON готов: {len(rows)} сообщений → {len(out_paths)} файл(ов)")
        return out_paths

    # ── Вспомогательные ───────────────────────────────────────────────

    @staticmethod
    def _make_record(row, stt_text: Optional[str]) -> dict:
        return {
            "message_id": row[_COL_MESSAGE_ID],
            "date":       row[_COL_DATE] or None,
            "sender_id":  row[_COL_USER_ID],
            "username":   row[_COL_USERNAME] or None,
            "text":       row[_COL_TEXT] or None,
            "media_path": row[_COL_MEDIA_PATH] or None,
            "stt_text":   stt_text,
        }

    def _write_json(self, path: str, records: List[dict], log: _LogCallback) -> None:
        log(f"💾 Записываю JSON: {os.path.basename(path)} ({len(records)} сообщений)")
        with open(path, "w", encoding="utf-8") as fh:
            json.dump(records, fh, ensure_ascii=False, indent=2, default=str)
        logger.info("JsonGenerator: saved %d records → %s", len(records), path)


# ==============================================================================
# MarkdownGenerator
# ==============================================================================

class MarkdownGenerator:
    """
    Генерирует Markdown-архив переписки из SQLite.

    Формат каждого сообщения:

        **[YYYY-MM-DD HH:MM] Имя Автора:**
        Текст сообщения

        *(STT: текст расшифровки)*   <- только если есть STT

        ---

    При ai_split=True файл разбивается на части по ai_split_chunk_words слов.

    Нет Qt-зависимостей. Нет Telethon-зависимостей. Только stdlib + DBManager.
    """

    def __init__(self, db: DBManager, output_dir: str = "output") -> None:
        self._db         = db
        self._output_dir = output_dir

    def generate(
        self,
        chat_id:              int,
        chat_title:           str,
        *,
        topic_id:             Optional[int]  = None,      # ← задача 2
        user_id:              Optional[int]  = None,
        include_comments:     bool           = False,
        ai_split:             bool           = False,
        period_label:         str,
        ai_split_chunk_words: int            = 300_000,    # ← задача 4
        date_from:            Optional[str]  = None,
        date_to:              Optional[str]  = None,
        log:                  _LogCallback   = lambda _: None,
    ) -> List[str]:
        """
        Основная точка входа. Строит Markdown и сохраняет на диск.

        Имя файла:
            <chat>_topic{N}_{period}_history.md   (с topic_id)
            <chat>_{period}_history.md            (без topic_id)

        Returns:
            Список абсолютных путей к созданным .md файлам.

        Raises:
            EmptyDataError: нет сообщений для экспорта.
            OSError: ошибка записи файла.
        """
        os.makedirs(self._output_dir, exist_ok=True)

        log("📋 Загружаю сообщения из БД для Markdown-экспорта...")
        rows = self._db.get_messages(
            chat_id,
            topic_id         = topic_id,
            user_id          = user_id,
            include_comments = include_comments,
            date_from        = date_from,
            date_to          = date_to,
        )

        if not rows:
            raise EmptyDataError(f"Нет сообщений для чата {chat_id}")

        log(f"📊 Строк получено: {len(rows)}")

        stt_map:    dict[int, str] = self._db.get_transcriptions_for_chat(chat_id)
        safe_title  = sanitize_filename(chat_title)
        topic_sfx   = _topic_suffix(topic_id)           # ← задача 2
        base_name   = f"{safe_title}{topic_sfx}_{period_label}_history"

        header = f"# {chat_title}\n\n"

        # ── Без разбивки: один файл ────────────────────────────────────
        if not ai_split:
            lines: List[str] = [header]
            for row in rows:
                lines.append(self._format_message(row, stt_map.get(row[_COL_MESSAGE_ID])))
            out_path = os.path.join(self._output_dir, f"{base_name}.md")
            self._write_md(out_path, lines, log)
            return [out_path]

        # ── С разбивкой по ai_split_chunk_words слов ───────────────────
        out_paths: List[str] = []
        chunk:     List[str] = [header]
        words      = 0
        part       = 1

        for row in rows:
            msg_id  = row[_COL_MESSAGE_ID]
            stt     = stt_map.get(msg_id)
            block   = self._format_message(row, stt)
            chunk.append(block)
            words += (
                _word_count(row[_COL_TEXT])
                + _word_count(stt)
            )

            if words >= ai_split_chunk_words:               # ← задача 4
                path = os.path.join(
                    self._output_dir,
                    f"{base_name}_part_{part}.md",
                )
                self._write_md(path, chunk, log)
                out_paths.append(path)
                chunk = [header]
                words = 0
                part += 1

        if len(chunk) > 1:   # есть сообщения (не только header)
            path = os.path.join(
                self._output_dir,
                f"{base_name}_part_{part}.md",
            )
            self._write_md(path, chunk, log)
            out_paths.append(path)

        log(f"✅ Markdown готов: {len(rows)} сообщений → {len(out_paths)} файл(ов)")
        return out_paths

    # ── Вспомогательные ───────────────────────────────────────────────

    @staticmethod
    def _format_message(row, stt_text: Optional[str]) -> str:
        """Форматирует одно сообщение в Markdown-блок."""
        raw_date = row[_COL_DATE] or ""
        date_str = raw_date[:16].replace("T", " ") if raw_date else "—"
        author   = row[_COL_USERNAME] or f"id:{row[_COL_USER_ID]}" or "Неизвестно"
        text     = (row[_COL_TEXT] or "").strip()

        lines = [f"**[{date_str}] {author}:**"]
        if text:
            lines.append(text)
        if stt_text:
            lines.append(f"\n*(STT: {stt_text.strip()})*")
        lines.append("\n---\n")
        return "\n".join(lines) + "\n"

    def _write_md(self, path: str, lines: List[str], log: _LogCallback) -> None:
        log(f"💾 Записываю MD: {os.path.basename(path)}")
        with open(path, "w", encoding="utf-8") as fh:
            fh.write("".join(lines))
        logger.info("MarkdownGenerator: saved → %s", path)


# ==============================================================================
# HtmlGenerator  — задача 6
# ==============================================================================

_HTML_TEMPLATE = """\
<!DOCTYPE html>
<html lang="ru">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>{title}</title>
<style>
  *, *::before, *::after {{ box-sizing: border-box; }}
  body {{
    background: #0e1117;
    font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, Helvetica, sans-serif;
    padding: 20px 12px;
    color: #e7edf5;
    margin: 0;
  }}
  .container {{ max-width: 800px; margin: 0 auto; }}
  .header {{
    margin-bottom: 24px;
    padding-bottom: 16px;
    border-bottom: 1px solid #2a2f3f;
    position: sticky;
    top: 0;
    background: #0e1117;
    z-index: 10;
    padding-top: 8px;
  }}
  .header h1 {{
    font-size: 1.5rem;
    font-weight: 600;
    background: linear-gradient(135deg, #FFB347, #FF8C42);
    background-clip: text;
    -webkit-background-clip: text;
    color: transparent;
  }}
  .header .stats {{ font-size: 0.75rem; color: #6c7a8e; margin-top: 6px; }}
  .message {{ margin-bottom: 4px; transition: background 0.1s; }}
  .message:target {{ background: rgba(255,149,0,0.15); border-radius: 12px; scroll-margin-top: 80px; }}
  .message.highlight {{ background: rgba(255,149,0,0.25); border-radius: 12px; }}
  .msg-row {{ display: flex; gap: 12px; padding: 10px 12px; border-radius: 12px; transition: background 0.1s; }}
  .msg-row:hover {{ background: #1a232e; }}
  .depth-0 {{ margin-left: 0; }}
  .depth-1 {{ margin-left: 32px; }}
  .depth-2 {{ margin-left: 56px; }}
  .depth-3 {{ margin-left: 80px; }}
  .avatar {{
    width: 36px; height: 36px; border-radius: 50%;
    display: flex; align-items: center; justify-content: center;
    font-weight: 600; font-size: 0.85rem; flex-shrink: 0;
    background: linear-gradient(135deg, #FF9500, #FF6B00);
    color: #0e1117;
  }}
  .content {{ flex: 1; min-width: 0; }}
  .msg-header {{ display: flex; align-items: baseline; flex-wrap: wrap; gap: 6px; margin-bottom: 4px; }}
  .author {{ font-weight: 600; font-size: 0.9rem; color: #FFA559; }}
  .date {{ font-size: 0.65rem; color: #6f7e93; }}
  .reply-badge {{ font-size: 0.65rem; color: #70b0ff; margin-left: 4px; }}
  .reply-badge a {{ color: #70b0ff; text-decoration: none; display: inline-flex; align-items: center; gap: 3px; }}
  .reply-badge a:hover {{ text-decoration: underline; }}
  .quote-preview {{
    font-size: 0.7rem; color: #9aaec9; margin-bottom: 6px;
    padding-left: 8px; border-left: 2px solid #FF9500; cursor: pointer;
  }}
  .msg-text {{ font-size: 0.88rem; line-height: 1.45; color: #e2eaf5; white-space: pre-wrap; word-break: break-word; }}
  .msg-text a {{ color: #70b0ff; text-decoration: none; border-bottom: 1px dotted #70b0ff; }}
  .msg-media {{
    margin-top: 8px; font-size: 0.75rem;
    background: #0e151e; padding: 4px 10px;
    border-radius: 12px; display: inline-block;
  }}
  .msg-media a {{ color: #60a0e0; text-decoration: none; }}
  .msg-media a:hover {{ text-decoration: underline; }}
  .msg-img {{ margin-top: 8px; max-width: 100%; max-height: 400px; border-radius: 8px; display: block; cursor: pointer; }}
  .msg-stt {{ margin-top: 6px; font-size: 0.82rem; color: #90a0b0; font-style: italic; padding-left: 4px; border-left: 2px solid #3a4a5a; }}
  .scroll-top {{
    position: fixed; bottom: 20px; right: 20px;
    background: #FF9500; color: #0e1117;
    border: none; border-radius: 40px; padding: 8px 16px;
    cursor: pointer; font-size: 0.8rem; font-weight: 600;
    box-shadow: 0 2px 12px rgba(0,0,0,0.4);
  }}
  .scroll-top:hover {{ background: #FFB347; }}
  .part-nav {{ text-align: center; padding: 16px 0; color: #607080; font-size: 0.82rem; }}
</style>
</head>
<body>
<div class="container">
  <div class="header">
    <h1>{title}</h1>
    <div class="stats">{total} сообщений</div>
  </div>
  <div id="messages-container">
{body}
  </div>
</div>
<button class="scroll-top" id="scrollTopBtn">↑ Наверх</button>
<script>
  document.getElementById('scrollTopBtn').addEventListener('click', function() {{
    window.scrollTo({{ top: 0, behavior: 'smooth' }});
  }});
  document.querySelectorAll('a[href^="#msg"]').forEach(function(link) {{
    link.addEventListener('click', function(e) {{
      var hash = this.getAttribute('href');
      if (hash && hash.startsWith('#')) {{
        e.preventDefault();
        var target = document.getElementById(hash.substring(1));
        if (target) {{
          target.scrollIntoView({{ behavior: 'smooth', block: 'center' }});
          target.classList.add('highlight');
          setTimeout(function() {{ target.classList.remove('highlight'); }}, 1000);
        }}
      }}
    }});
  }});
</script>
</body>
</html>
"""

_HTML_MSG = """\
<div class="message depth-{depth}" id="msg_{msg_id}">
  <div class="msg-row">
    <div class="avatar">{avatar_letter}</div>
    <div class="content">
      <div class="msg-header">
        <span class="author">{author}</span>
        <span class="date">{date}</span>
        {reply_badge}
      </div>
      {quote_block}{text_block}{media_block}{stt_block}
    </div>
  </div>
</div>
"""


class HtmlGenerator:
    """
    Генерирует HTML-архив переписки из SQLite.

    Структура выходного файла — тёмная тема, адаптивная вёрстка.
    При ai_split=True файл разбивается на части по ai_split_chunk_words слов.

    Нет Qt-зависимостей. Нет Telethon-зависимостей. Только stdlib + DBManager.
    """

    def __init__(self, db: DBManager, output_dir: str = "output") -> None:
        self._db         = db
        self._output_dir = output_dir

    def generate(
        self,
        chat_id:              int,
        chat_title:           str,
        *,
        topic_id:             Optional[int]  = None,
        user_id:              Optional[int]  = None,
        include_comments:     bool           = False,
        ai_split:             bool           = False,
        period_label:         str            = "fullchat",
        ai_split_chunk_words: int            = 300_000,
        date_from:            Optional[str]  = None,
        date_to:              Optional[str]  = None,
        log:                  _LogCallback   = lambda _: None,
    ) -> List[str]:
        """
        Основная точка входа. Строит HTML и сохраняет на диск.

        Имя файла:
            <chat>_topic{N}_{period}_history.html   (с topic_id)
            <chat>_{period}_history.html            (без topic_id)
            <chat>_topic{N}_{period}_history_part_{k}.html  (с ai_split)

        Returns:
            Список абсолютных путей к созданным .html файлам.

        Raises:
            EmptyDataError: нет сообщений для экспорта.
            OSError: ошибка записи файла.
        """
        os.makedirs(self._output_dir, exist_ok=True)

        log("📋 Загружаю сообщения из БД для HTML-экспорта...")
        rows = self._db.get_messages(
            chat_id,
            topic_id         = topic_id,
            user_id          = user_id,
            include_comments = include_comments,
            date_from        = date_from,
            date_to          = date_to,
        )

        if not rows:
            raise EmptyDataError(f"Нет сообщений для чата {chat_id}")

        log(f"📊 Строк получено: {len(rows)}")

        stt_map:    dict[int, str] = self._db.get_transcriptions_for_chat(chat_id)
        safe_title  = sanitize_filename(chat_title)
        topic_sfx   = _topic_suffix(topic_id)
        base_name   = f"{safe_title}{topic_sfx}_{period_label}_history"
        h_title     = html_lib.escape(chat_title)

        total    = len(rows)
        row_dict = {row[_COL_MESSAGE_ID]: row for row in rows}

        # ── Без разбивки: один файл ────────────────────────────────────
        if not ai_split:
            blocks: List[str] = []
            for row in rows:
                blocks.append(self._format_message(row, stt_map.get(row[_COL_MESSAGE_ID]), row_dict))
            out_path = os.path.join(self._output_dir, f"{base_name}.html")
            self._write_html(out_path, h_title, blocks, total, log)
            return [out_path]

        # ── С разбивкой по ai_split_chunk_words слов ───────────────────
        out_paths: List[str] = []
        chunk:     List[str] = []
        words      = 0
        part       = 1

        for row in rows:
            msg_id = row[_COL_MESSAGE_ID]
            stt    = stt_map.get(msg_id)
            block  = self._format_message(row, stt, row_dict)
            chunk.append(block)
            words += _word_count(row[_COL_TEXT]) + _word_count(stt)

            if words >= ai_split_chunk_words:
                path = os.path.join(self._output_dir, f"{base_name}_part_{part}.html")
                nav = f'<div class="part-nav">Часть {part} · {len(chunk)} сообщений</div>'
                self._write_html(path, f"{h_title} — часть {part}", chunk + [nav], total, log)
                out_paths.append(path)
                chunk = []
                words = 0
                part += 1

        if chunk:
            path = os.path.join(self._output_dir, f"{base_name}_part_{part}.html")
            nav = f'<div class="part-nav">Часть {part} · {len(chunk)} сообщений</div>'
            self._write_html(path, f"{h_title} — часть {part}", chunk + [nav], total, log)
            out_paths.append(path)

        log(f"✅ HTML готов: {total} сообщений → {len(out_paths)} файл(ов)")
        return out_paths

    # ── Вспомогательные ───────────────────────────────────────────────

    @staticmethod
    def _format_message(row, stt_text: Optional[str], row_dict: dict) -> str:
        """Форматирует одно сообщение в HTML-блок по структуре макета."""
        msg_id   = row[_COL_MESSAGE_ID]
        raw_date = row[_COL_DATE] or ""
        date_str = raw_date[:16].replace("T", " ") if raw_date else "—"
        author   = row[_COL_USERNAME] or f"id:{row[_COL_USER_ID]}" or "Неизвестно"
        text     = (row[_COL_TEXT] or "").strip()
        reply_to = row[_COL_REPLY_TO]

        # Первая буква автора для аватара
        avatar_letter = html_lib.escape(author)[0].upper() if author else "?"

        # Глубина: 1 если есть ответ на известное сообщение, иначе 0
        depth = 1 if (reply_to and reply_to in row_dict) else 0

        # Бейдж ответа
        reply_badge = ""
        if reply_to and reply_to in row_dict:
            reply_badge = (
                f'<span class="reply-badge">'
                f'<a href="#msg_{reply_to}">↪️ ответ</a></span>'
            )

        # Блок цитаты (превью сообщения, на которое отвечают)
        quote_block = ""
        if reply_to and reply_to in row_dict:
            ref = row_dict[reply_to]
            ref_author  = ref[_COL_USERNAME] or f"id:{ref[_COL_USER_ID]}" or "?"
            ref_text    = (ref[_COL_TEXT] or "").strip()
            preview     = ref_text[:70] + ("…" if len(ref_text) > 70 else "")
            quote_block = (
                f'<div class="quote-preview" onclick="'
                f"var t=document.getElementById('msg_{reply_to}');"
                f"if(t){{t.scrollIntoView({{behavior:'smooth',block:'center'}});"
                f"t.classList.add('highlight');"
                f"setTimeout(function(){{t.classList.remove('highlight')}},1000)}}"
                f'">'
                f'💬 {html_lib.escape(ref_author)}: {html_lib.escape(preview)}'
                f'</div>\n      '
            )

        # Текст с авто-ссылками
        text_block = ""
        if text:
            escaped = html_lib.escape(text)
            linked  = re.sub(
                r'(https?://[^\s"\'<>)]+)',
                r'<a href="\1" target="_blank">\1</a>',
                escaped,
            )
            text_block = f'<div class="msg-text">{linked}</div>\n      '

        # Медиа: inline-превью для изображений, ссылка для остальных
        media_block = ""
        media_path  = row[_COL_MEDIA_PATH]
        if media_path:
            fname = os.path.basename(media_path)
            if os.path.exists(media_path):
                uri = Path(os.path.abspath(media_path)).as_uri()
                if is_image_path(media_path):
                    media_block = (
                        f'<img class="msg-img" src="{uri}" alt="{html_lib.escape(fname)}"'
                        f' onclick="window.open(\'{uri}\',\'_blank\')">\n      '
                        f'<div class="msg-media">🖼 '
                        f'<a href="{uri}" target="_blank">{html_lib.escape(fname)}</a>'
                        f'</div>\n      '
                    )
                else:
                    media_block = (
                        f'<div class="msg-media">📎 '
                        f'<a href="{uri}" target="_blank">{html_lib.escape(fname)}</a>'
                        f'</div>\n      '
                    )
            else:
                media_block = (
                    f'<div class="msg-media">📎 [недоступен] '
                    f'{html_lib.escape(fname)}</div>\n      '
                )

        # STT — курсив с левой полосой
        stt_block = ""
        if stt_text:
            stt_block = (
                f'<div class="msg-stt">🎙 {html_lib.escape(stt_text.strip())}</div>\n      '
            )

        return _HTML_MSG.format(
            depth         = depth,
            msg_id        = msg_id,
            avatar_letter = avatar_letter,
            author        = html_lib.escape(author),
            date          = html_lib.escape(date_str),
            reply_badge   = reply_badge,
            quote_block   = quote_block,
            text_block    = text_block,
            media_block   = media_block,
            stt_block     = stt_block,
        )

    def _write_html(
        self,
        path:   str,
        title:  str,
        blocks: List[str],
        total:  int,
        log:    _LogCallback,
    ) -> None:
        log(f"💾 Записываю HTML: {os.path.basename(path)} ({len(blocks)} блоков)")
        content = _HTML_TEMPLATE.format(title=title, body="".join(blocks), total=total)
        with open(path, "w", encoding="utf-8") as fh:
            fh.write(content)
        logger.info("HtmlGenerator: saved → %s", path)
