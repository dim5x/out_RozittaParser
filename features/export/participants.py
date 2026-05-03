"""
features/export/participants.py — Экспорт списка участников.

Поддерживает формат DOCX:
  - DOCX (рекомендуемый): таблица с активными ссылками на профили Telegram.

Нет импортов Qt. Нет Telethon. Только stdlib + python-docx + core.utils.

Публичный API:
    export_participants_docx(users, chat_title, output_dir, counts) → str
        Создаёт DOCX с таблицей и кликабельными ссылками на профили.
"""

from datetime import datetime
import logging
import os

from core.utils import sanitize_filename

logger = logging.getLogger(__name__)

def export_participants_docx(
        users: list[dict],
        chat_title: str,
        output_dir: str,
        # counts,
) -> str:
    """
    Создаёт DOCX-файл со списком участников.

    Таблица содержит:
      - #  (порядковый номер)
      - Имя (отображаемое, жирное)
      - @username (кликабельная ссылка tg://user?id=... или https://t.me/username)
      - Сообщений (если есть данные)

    Строки отсортированы по убыванию активности.

    Returns:
        Абсолютный путь к созданному .docx файлу.
    """

    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from features.export import xml_magic

    os.makedirs(output_dir, exist_ok=True)

    now = datetime.now()
    date_str = now.strftime("%Y-%m-%d %H:%M")
    file_date = now.strftime("%Y-%m-%d_%H-%M")
    safe_title = sanitize_filename(chat_title)
    filename = f"{safe_title}_participants_{file_date}.docx"
    filepath = os.path.join(output_dir, filename)

    has_counts = any(u['message_count'] > 0 for u in users)
    sorted_users = sorted(users, key=lambda x: x['message_count'], reverse=True)
    total_msgs = sum(u["message_count"] for u in sorted_users)

    doc = Document()

    # ── Стиль документа ───────────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ── Заголовок ─────────────────────────────────────────────────────────
    title_p = doc.add_heading(f"Участники: {chat_title}", level=1)
    title_p.runs[0].font.color.rgb = RGBColor(0xFF, 0x6B, 0xC9)  # ACCENT_PINK

    # ── Метаданные ────────────────────────────────────────────────────────
    meta_p = doc.add_paragraph()
    meta_p.add_run(f"Дата выгрузки: ").bold = False
    meta_p.add_run(date_str).bold = True
    meta_p.add_run(f"   •   Участников: ").bold = False
    meta_p.add_run(str(len(sorted_users))).bold = True
    if total_msgs:
        meta_p.add_run("   •   Сообщений в архиве: ").bold = False
        meta_p.add_run(f"{total_msgs:,}").bold = True
    meta_p.runs[0].font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    doc.add_paragraph()  # пустая строка

    # ── Таблица ───────────────────────────────────────────────────────────
    col_count = 4 if has_counts else 3
    table = doc.add_table(rows=1, cols=col_count)
    table.style = "Table Grid"

    # Заголовки
    hdr_cells = table.rows[0].cells
    headers = ["#", "Имя", "Ссылка / @username"]
    if has_counts:
        headers.append("Сообщений")

    for i, text in enumerate(headers):
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0x95, 0x00)  # ACCENT_ORANGE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in (0, col_count - 1) else WD_ALIGN_PARAGRAPH.LEFT

    # Ширина столбцов
    col_widths = [Cm(1.0), Cm(5.5), Cm(7.5), Cm(2.5)][:col_count]
    for i, cell in enumerate(hdr_cells):
        cell.width = col_widths[i]

    # Строки участников
    for idx, user in enumerate(sorted_users, start=1):
        uid = user.get("id", 0)
        name = user.get("name", "")
        username = user.get("username", "")
        count = user.get("message_count", 0)

        row_cells = table.add_row().cells

        # #
        p_num = row_cells[0].paragraphs[0]
        p_num.add_run(str(idx))
        p_num.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Имя
        p_name = row_cells[1].paragraphs[0]
        r = p_name.add_run(name)
        r.bold = True

        # Ссылка
        p_link = row_cells[2].paragraphs[0]
        if username:
            # Ссылка по username: открывается в браузере и в Telegram
            url = f"https://t.me/{username}"
            link_text = f"@{username}"
        else:
            # Ссылка по ID: открывает чат в Telegram Desktop/Mobile
            url = f"tg://user?id={uid}"
            link_text = f"tg://user?id={uid}"

        xml_magic.add_external_hyperlink(p_link, url, link_text)

        # Сообщений
        if has_counts:
            p_cnt = row_cells[3].paragraphs[0]
            p_cnt.add_run(f"{count:,}" if count else "—")
            p_cnt.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Ширина столбцов — устанавливаем для каждой строки
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = col_widths[i]

    doc.save(filepath)
    logger.info("participants: docx exported %d users → %s", len(sorted_users), filepath)
    return filepath
