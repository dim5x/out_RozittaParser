"""
features/export/xml_magic.py — Низкоуровневые XML-операции для DOCX.

Содержит функции, которые python-docx не поддерживает «из коробки»:
  - add_bookmark()              — невидимая закладка (якорь) на параграф
  - add_internal_hyperlink()    — ссылка на закладку внутри того же документа
  - add_external_hyperlink()    — кликабельная внешняя ссылка (http://, file://)
  - write_text_with_links()     — авто-разбивка текста на plain text + URL-ссылки

Все функции работают на уровне OxmlElement (lxml) и не зависят ни от Qt,
ни от Telethon, ни от базы данных.

Как это работает в Word:
    Bookmark:  <w:bookmarkStart w:id="N" w:name="msg_123"/>...<w:bookmarkEnd w:id="N"/>
    Internal:  <w:hyperlink w:anchor="msg_123"><w:r>...<w:t>текст</w:t></w:r></w:hyperlink>
    External:  <w:hyperlink r:id="rIdN"><w:r>...<w:t>текст</w:t></w:r></w:hyperlink>
               где rIdN — relationship_id в part/_rels/*.xml.rels

Проблема с w:id в закладках:
    Word требует уникальные w:id в документе. Если добавлять несколько закладок
    с одним id=0, документ откроется, но навигация сломается. Используем
    глобальный счётчик _bookmark_id_counter (сбрасывается функцией reset_counter()).
    reset_counter() вызывается в начале каждого нового Document() в generator.py.
"""

from __future__ import annotations

import re
import logging

import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

logger = logging.getLogger(__name__)

# Глобальный счётчик ID для закладок (уникален в пределах одного Document)
_bookmark_counter: int = 0


def reset_counter() -> None:
    """
    Сбрасывает счётчик закладок в 0.

    ОБЯЗАТЕЛЬНО вызывать перед созданием каждого нового Document(),
    иначе w:id будут продолжать расти из предыдущего документа
    (функционально не критично, но нарушает спецификацию OOXML).

    Example:
        xml_magic.reset_counter()
        doc = Document()
        ...
    """

    global _bookmark_counter
    _bookmark_counter = 0


# ==============================================================================
# Закладки
# ==============================================================================

def add_bookmark(paragraph: Paragraph, bookmark_name: str) -> None:
    """
    Добавляет невидимую закладку (anchor) в конец параграфа.

    Закладка — это пустой XML-элемент, к которому можно создать
    внутреннюю ссылку через add_internal_hyperlink().

    Соглашение об именовании: "msg_{message_id}", например "msg_42".

    Args:
        paragraph:     Параграф python-docx, в который вставляется закладка.
        bookmark_name: Уникальное имя закладки (ASCII, без пробелов).

    Example:
        p = doc.add_paragraph()
        add_bookmark(p, "msg_42")
    """

    global _bookmark_counter
    bm_id = str(_bookmark_counter)
    _bookmark_counter += 1

    tag = paragraph._p

    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"),   bm_id)
    start.set(qn("w:name"), bookmark_name)
    tag.append(start)

    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), bm_id)
    tag.append(end)

    logger.debug("xml_magic: add_bookmark '%s' id=%s", bookmark_name, bm_id)


# ==============================================================================
# Внутренние ссылки
# ==============================================================================

def add_internal_hyperlink(
    paragraph:   Paragraph,
    target_id:   int,
    text:        str,
    color_hex:   str = "0563C1",
) -> None:
    """
    Добавляет кликабельную ссылку на закладку "msg_{target_id}" внутри документа.

    Ссылка оформляется синим цветом с подчёркиванием (стандарт Word).
    При нажатии Word переходит к параграфу с соответствующей закладкой.

    Args:
        paragraph:  Параграф, в конец которого добавляется ссылка.
        target_id:  ID сообщения-цели (будет искать закладку "msg_{target_id}").
        text:       Текст ссылки (например, "сообщение #42").
        color_hex:  Цвет текста ссылки в HEX без '#' (по умолчанию синий Word).

    Example:
        p = doc.add_paragraph()
        p.add_run("↩️ В ответ на: ")
        add_internal_hyperlink(p, reply_to_msg_id, f"сообщение #{reply_to_msg_id}")
    """

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), f"msg_{target_id}")

    run = OxmlElement("w:r")

    rPr = OxmlElement("w:rPr")
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), color_hex)
    rPr.append(color_el)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rPr.append(underline)
    run.append(rPr)

    t = OxmlElement("w:t")
    t.text = text
    run.append(t)

    hyperlink.append(run)
    paragraph._p.append(hyperlink)

    logger.debug("xml_magic: internal_link → msg_%s '%s'", target_id, text)


# ==============================================================================
# Внешние ссылки
# ==============================================================================

def add_external_hyperlink(
    paragraph: Paragraph,
    url:       str,
    text:      str,
    color_hex: str = "0563C1",
) -> None:
    """
    Добавляет кликабельную внешнюю ссылку (http://, https://, file://).

    Ссылка регистрируется в relationships части документа (part.relate_to),
    что обязательно для корректной работы в Word и LibreOffice.

    Args:
        paragraph:  Параграф, в конец которого добавляется ссылка.
        url:        Полный URL (например, "https://t.me/channel/123" или
                    "file:///C:/output/media/photo.jpg").
        text:       Текст ссылки.
        color_hex:  Цвет текста ссылки в HEX без '#'.

    Example:
        from pathlib import Path
        media_uri = Path(abs_path).as_uri()
        p = doc.add_paragraph("📎 Медиафайл: ")
        add_external_hyperlink(p, media_uri, "photo.jpg")
    """

    part = paragraph.part
    r_id = part.relate_to(
        url,
        docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK,
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")

    rPr = OxmlElement("w:rPr")
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), color_hex)
    rPr.append(color_el)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rPr.append(underline)
    run.append(rPr)

    t = OxmlElement("w:t")
    t.text = text
    run.append(t)

    hyperlink.append(run)
    paragraph._p.append(hyperlink)

    logger.debug("xml_magic: external_link '%s' → %s", text, url[:60])


# ==============================================================================
# Авто-разбивка текста на plain text + ссылки
# ==============================================================================

# Regex для поиска URL в тексте
_URL_RE = re.compile(
    r"(https?://[^\s\"'<>)]+)",
    re.IGNORECASE,
)


def write_text_with_links(paragraph: Paragraph, text: str) -> None:
    """
    Добавляет текст в параграф, автоматически превращая URL в кликабельные ссылки.

    Алгоритм:
        1. Разбивает текст по regex _URL_RE
        2. Нечётные части (индексы 0, 2, 4...) — plain text → add_run()
        3. Чётные части (индексы 1, 3, 5...) — URL → add_external_hyperlink()

    Args:
        paragraph: Параграф, в который пишем текст.
        text:      Строка с произвольным текстом (может содержать URL).

    Example:
        p = doc.add_paragraph()
        write_text_with_links(p, "Смотрите: https://example.com и обычный текст")
        # → "Смотрите: " (plain) + "https://example.com" (hyperlink) + " и обычный текст" (plain)
    """

    if not text:
        return

    parts = _URL_RE.split(text)

    for i, part in enumerate(parts):
        if not part:
            continue

        if i % 2 == 1:
            # URL-часть — делаем ссылкой
            add_external_hyperlink(paragraph, part, part)
        else:
            # Обычный текст
            paragraph.add_run(part)
