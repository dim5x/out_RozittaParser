"""
tests/test_e2e/test_full_cycle.py

E2E: полный цикл insert msgs + transcriptions → generate DOCX/JSON/MD/HTML.
"""
import json
import re

import pytest
from docx import Document

from core.database import DBManager
from features.export.generator import (
    DocxGenerator,
    HtmlGenerator,
    JsonGenerator,
    MarkdownGenerator,
)


CHAT_ID = -100200


def _msg(i, chat_id=CHAT_ID, text=None, file_type=None, user_id=None):
    return {
        "chat_id": chat_id,
        "message_id": i,
        "topic_id": None,
        "user_id": user_id or (i % 5) + 1,
        "username": f"user_{user_id or (i % 5) + 1}",
        "date": f"2025-01-{1 + (i // 24):02d}T{i % 24:02d}:00:00",
        "text": text if text is not None else (f"Message number {i}" if i % 3 != 0 else ""),
        "media_path": None,
        "file_type": file_type or ("voice" if i % 7 == 0 else None),
        "file_size": None,
        "reply_to_msg_id": None,
        "post_id": None,
        "is_comment": 0,
        "from_linked_group": 0,
    }


@pytest.fixture
def populated_db(tmp_path):
    path = str(tmp_path / "archive.db")
    db = DBManager(path)
    msgs = [_msg(i) for i in range(1, 51)]
    db.insert_messages_batch(msgs)

    for i in range(1, 51):
        if i % 7 == 0:
            db.insert_transcription(i, CHAT_ID, f"Транскрипция {i}", "base")

    return db, str(tmp_path)


class TestDocxFullCycle:
    def test_generates_valid_docx(self, populated_db):
        db, out = populated_db
        gen = DocxGenerator(db=db, output_dir=out)
        files = gen.generate(
            chat_id=CHAT_ID, chat_title="Full Cycle Test",
            split_mode="none", period_label="fullchat",
        )
        assert len(files) == 1
        assert files[0].endswith(".docx")
        doc = Document(files[0])
        assert len(doc.paragraphs) > 0

    def test_docx_contains_stt(self, populated_db):
        db, out = populated_db
        gen = DocxGenerator(db=db, output_dir=out)
        files = gen.generate(
            chat_id=CHAT_ID, chat_title="Full Cycle Test",
            split_mode="none", period_label="fullchat",
        )
        doc = Document(files[0])
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Распознанная речь" in full_text


class TestJsonFullCycle:
    def test_generates_valid_json(self, populated_db):
        db, out = populated_db
        gen = JsonGenerator(db=db, output_dir=out)
        files = gen.generate(
            chat_id=CHAT_ID, chat_title="Full Cycle Test", period_label="fullchat",
        )
        assert len(files) == 1
        with open(files[0], encoding="utf-8") as f:
            data = json.load(f)
        assert len(data) == 50

    def test_json_includes_stt(self, populated_db):
        db, out = populated_db
        gen = JsonGenerator(db=db, output_dir=out)
        files = gen.generate(
            chat_id=CHAT_ID, chat_title="Full Cycle Test", period_label="fullchat",
        )
        with open(files[0], encoding="utf-8") as f:
            data = json.load(f)
        stt_records = [r for r in data if r.get("stt_text")]
        assert len(stt_records) > 0

    def test_json_fields_present(self, populated_db):
        db, out = populated_db
        gen = JsonGenerator(db=db, output_dir=out)
        files = gen.generate(
            chat_id=CHAT_ID, chat_title="Full Cycle Test", period_label="fullchat",
        )
        with open(files[0], encoding="utf-8") as f:
            data = json.load(f)
        required = {"message_id", "sender_id", "username", "date", "text"}
        for record in data:
            assert required <= set(record.keys())


class TestMdFullCycle:
    def test_generates_md(self, populated_db):
        db, out = populated_db
        gen = MarkdownGenerator(db=db, output_dir=out)
        files = gen.generate(
            chat_id=CHAT_ID, chat_title="Full Cycle Test", period_label="fullchat",
        )
        assert len(files) == 1
        assert files[0].endswith(".md")

    def test_md_includes_stt(self, populated_db):
        db, out = populated_db
        gen = MarkdownGenerator(db=db, output_dir=out)
        files = gen.generate(
            chat_id=CHAT_ID, chat_title="Full Cycle Test", period_label="fullchat",
        )
        content = open(files[0], encoding="utf-8").read()
        assert "STT:" in content

    def test_md_message_count(self, populated_db):
        db, out = populated_db
        gen = MarkdownGenerator(db=db, output_dir=out)
        files = gen.generate(
            chat_id=CHAT_ID, chat_title="Full Cycle Test", period_label="fullchat",
        )
        content = open(files[0], encoding="utf-8").read()
        blocks = re.findall(r"\*\[\d{4}-\d{2}-\d{2}", content)
        assert len(blocks) == 50


class TestHtmlFullCycle:
    def test_generates_html(self, populated_db):
        db, out = populated_db
        gen = HtmlGenerator(db=db, output_dir=out)
        files = gen.generate(
            chat_id=CHAT_ID, chat_title="Full Cycle Test", period_label="fullchat",
        )
        assert len(files) == 1
        assert files[0].endswith(".html")

    def test_html_includes_stt(self, populated_db):
        db, out = populated_db
        gen = HtmlGenerator(db=db, output_dir=out)
        files = gen.generate(
            chat_id=CHAT_ID, chat_title="Full Cycle Test", period_label="fullchat",
        )
        html = open(files[0], encoding="utf-8").read()
        assert "msg-stt" in html


class TestAllFormatsConsistency:
    def test_json_and_md_same_count(self, populated_db):
        db, out = populated_db

        jfiles = JsonGenerator(db=db, output_dir=out).generate(
            chat_id=CHAT_ID, chat_title="T", period_label="fullchat",
        )
        with open(jfiles[0], encoding="utf-8") as f:
            jdata = json.load(f)

        mfiles = MarkdownGenerator(db=db, output_dir=out).generate(
            chat_id=CHAT_ID, chat_title="T", period_label="fullchat",
        )
        md = open(mfiles[0], encoding="utf-8").read()
        md_count = len(re.findall(r"\*\[\d{4}-\d{2}-\d{2}", md))

        assert len(jdata) == md_count == 50

    def test_html_and_json_same_count(self, populated_db):
        db, out = populated_db

        jfiles = JsonGenerator(db=db, output_dir=out).generate(
            chat_id=CHAT_ID, chat_title="T", period_label="fullchat",
        )
        with open(jfiles[0], encoding="utf-8") as f:
            jdata = json.load(f)

        hfiles = HtmlGenerator(db=db, output_dir=out).generate(
            chat_id=CHAT_ID, chat_title="T", period_label="fullchat",
        )
        html = open(hfiles[0], encoding="utf-8").read()
        html_count = html.count('id="msg_')

        assert len(jdata) == html_count == 50


class TestAiSplitEndToEnd:
    def _make_split_msgs(self, n=100, chat_id=-100):
        return [_msg(i, chat_id=chat_id, text=f"{'word ' * 20}", file_type=None)
                for i in range(n)]

    def test_json_ai_split(self, tmp_path):
        db = DBManager(str(tmp_path / "split.db"))
        db.insert_messages_batch(self._make_split_msgs(100))

        gen = JsonGenerator(db=db, output_dir=str(tmp_path))
        files = gen.generate(
            chat_id=-100, chat_title="T", period_label="full",
            ai_split=True, ai_split_chunk_words=30,
        )
        assert len(files) >= 2
        total_records = 0
        for f in files:
            with open(f, encoding="utf-8") as fh:
                total_records += len(json.load(fh))
        assert total_records == 100

    def test_md_ai_split(self, tmp_path):
        db = DBManager(str(tmp_path / "split.db"))
        db.insert_messages_batch(self._make_split_msgs(100))

        gen = MarkdownGenerator(db=db, output_dir=str(tmp_path))
        files = gen.generate(
            chat_id=-100, chat_title="T", period_label="full",
            ai_split=True, ai_split_chunk_words=30,
        )
        assert len(files) >= 2

    def test_html_ai_split(self, tmp_path):
        db = DBManager(str(tmp_path / "split.db"))
        db.insert_messages_batch(self._make_split_msgs(100))

        gen = HtmlGenerator(db=db, output_dir=str(tmp_path))
        files = gen.generate(
            chat_id=-100, chat_title="T", period_label="full",
            ai_split=True, ai_split_chunk_words=30,
        )
        assert len(files) >= 2
        for f in files:
            html = open(f, encoding="utf-8").read()
            assert "<!DOCTYPE html>" in html
