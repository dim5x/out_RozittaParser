"""
tests/test_features/test_export_docx.py

Тесты: DocxGenerator edge cases — split modes, STT, merge groups, invalid inputs.
"""
import os
import pytest
from docx import Document
from core.database import DBManager
from core.exceptions import DocxGenerationError, EmptyDataError
from features.export.generator import DocxGenerator


def _msg(chat_id=-100, msg_id=1, user_id=1, username="user1",
         date="2025-01-01T10:00:00", text="hello", **kw):
    """Фабрика сообщения со всеми обязательными полями."""
    return {
        "chat_id": chat_id, "message_id": msg_id, "user_id": user_id,
        "username": username, "date": date, "text": text,
        "topic_id": None, "media_path": None, "file_type": None,
        "file_size": None, "reply_to_msg_id": None, "post_id": None,
        "is_comment": 0, "from_linked_group": 0,
        **kw,
    }


def _make_db(tmp_path, msgs):
    path = str(tmp_path / "test.db")
    db = DBManager(path)
    db.insert_messages_batch(msgs)
    return db


def _make_db_with_dates(tmp_path, dates):
    msgs = [_msg(msg_id=i, date=d, text=f"Message {i}", chat_id=-100123)
            for i, d in enumerate(dates, 1)]
    return _make_db(tmp_path, msgs)


class TestDocxSplitByDay:
    def test_two_days_two_files(self, tmp_path):
        db = _make_db_with_dates(tmp_path, [
            "2025-01-01T10:00:00",
            "2025-01-01T14:00:00",
            "2025-01-02T09:00:00",
        ])
        gen = DocxGenerator(db=db, output_dir=str(tmp_path))
        files = gen.generate(
            chat_id=-100123, chat_title="Test", split_mode="day",
            period_label="fullchat",
        )
        assert len(files) == 2

    def test_one_day_one_file(self, tmp_path):
        db = _make_db_with_dates(tmp_path, [
            "2025-01-01T10:00:00",
            "2025-01-01T14:00:00",
        ])
        gen = DocxGenerator(db=db, output_dir=str(tmp_path))
        files = gen.generate(
            chat_id=-100123, chat_title="Test", split_mode="day",
            period_label="fullchat",
        )
        assert len(files) == 1

    def test_all_files_exist(self, tmp_path):
        db = _make_db_with_dates(tmp_path, [
            "2025-01-01T10:00:00",
            "2025-01-02T10:00:00",
        ])
        gen = DocxGenerator(db=db, output_dir=str(tmp_path))
        files = gen.generate(
            chat_id=-100123, chat_title="Test", split_mode="day",
            period_label="fullchat",
        )
        for f in files:
            assert os.path.exists(f)
            assert f.endswith(".docx")


class TestDocxSplitByMonth:
    def test_two_months_two_files(self, tmp_path):
        db = _make_db_with_dates(tmp_path, [
            "2025-01-15T10:00:00",
            "2025-02-20T10:00:00",
        ])
        gen = DocxGenerator(db=db, output_dir=str(tmp_path))
        files = gen.generate(
            chat_id=-100123, chat_title="Test", split_mode="month",
            period_label="fullchat",
        )
        assert len(files) == 2

    def test_same_month_one_file(self, tmp_path):
        db = _make_db_with_dates(tmp_path, [
            "2025-01-10T10:00:00",
            "2025-01-20T14:00:00",
        ])
        gen = DocxGenerator(db=db, output_dir=str(tmp_path))
        files = gen.generate(
            chat_id=-100123, chat_title="Test", split_mode="month",
            period_label="fullchat",
        )
        assert len(files) == 1


class TestDocxSplitInvalid:
    def test_invalid_split_mode_raises(self, tmp_path):
        db = _make_db(tmp_path, [_msg(chat_id=-100)])
        gen = DocxGenerator(db=db, output_dir=str(tmp_path))
        with pytest.raises(DocxGenerationError, match="split_mode"):
            gen.generate(chat_id=-100, chat_title="T", split_mode="bad_mode")

    def test_empty_db_raises(self, tmp_path):
        db = DBManager(str(tmp_path / "empty.db"))
        gen = DocxGenerator(db=db, output_dir=str(tmp_path))
        with pytest.raises(EmptyDataError):
            gen.generate(chat_id=-100, chat_title="T", split_mode="none")


class TestDocxStt:
    def test_stt_text_in_docx(self, tmp_path):
        db = DBManager(str(tmp_path / "test.db"))
        db.insert_messages_batch([_msg(text="", file_type="voice", media_path="/a.ogg")])
        db.insert_transcription(1, -100, "Привет мир", "base")
        gen = DocxGenerator(db=db, output_dir=str(tmp_path))
        files = gen.generate(chat_id=-100, chat_title="T", split_mode="none")
        assert len(files) == 1
        doc = Document(files[0])
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Привет мир" in full_text


class TestDocxMergeGroups:
    def test_merge_group_output(self, tmp_path):
        db = DBManager(str(tmp_path / "test.db"))
        db.insert_messages_batch([
            _msg(msg_id=1, username="A", text="Привет",
                 merge_group_id=10, merge_part_index=0),
            _msg(msg_id=2, username="A", text="мир",
                 date="2025-01-01T10:00:20",
                 merge_group_id=10, merge_part_index=1),
        ])
        gen = DocxGenerator(db=db, output_dir=str(tmp_path))
        files = gen.generate(chat_id=-100, chat_title="T", split_mode="none")
        doc = Document(files[0])
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Привет" in full_text
        assert "мир" in full_text


class TestDocxTopicSuffix:
    def test_filename_with_topic(self, tmp_path):
        db = DBManager(str(tmp_path / "test.db"))
        db.insert_messages_batch([_msg(text="hi", topic_id=7)])
        gen = DocxGenerator(db=db, output_dir=str(tmp_path))
        files = gen.generate(
            chat_id=-100, chat_title="T", split_mode="none",
            topic_id=7, period_label="fullchat",
        )
        assert len(files) == 1
        assert "_topic7" in files[0]

    def test_filename_without_topic(self, tmp_path):
        db = _make_db(tmp_path, [_msg(chat_id=-100)])
        gen = DocxGenerator(db=db, output_dir=str(tmp_path))
        files = gen.generate(
            chat_id=-100, chat_title="T", split_mode="none",
            period_label="fullchat",
        )
        assert "_topic" not in os.path.basename(files[0])


class TestDocxUnicode:
    def test_unicode_text_in_docx(self, tmp_path):
        db = _make_db(tmp_path, [_msg(text="Привет мир", username="Алексей")])
        gen = DocxGenerator(db=db, output_dir=str(tmp_path))
        files = gen.generate(chat_id=-100, chat_title="Тест", split_mode="none")
        doc = Document(files[0])
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Привет мир" in full_text


class TestDocxContentVerification:
    def test_message_text_in_docx(self, tmp_path):
        db = _make_db(tmp_path, [
            _msg(msg_id=1, username="Alice", text="First message"),
            _msg(msg_id=2, username="Bob", text="Second message", date="2025-01-01T10:05:00"),
        ])
        gen = DocxGenerator(db=db, output_dir=str(tmp_path))
        files = gen.generate(chat_id=-100, chat_title="T", split_mode="none")
        doc = Document(files[0])
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "First message" in full_text
        assert "Second message" in full_text
