"""
tests/test_features/test_participants.py

Тесты: export_participants_docx — создание DOCX с таблицей участников.
Использует реальный Document() для проверки структуры.
"""
import os

from docx import Document

from features.export.participants import export_participants_docx


def _user(id_, name="User", username="", message_count=0):
    return {
        "id": id_,
        "name": name,
        "username": username,
        "message_count": message_count,
    }


class TestParticipantsBasic:
    def test_creates_docx_file(self, tmp_path):
        path = export_participants_docx(
            [_user(1, "Alice", "alice", 10)],
            "TestChat", str(tmp_path),
        )
        assert os.path.exists(path)
        assert path.endswith(".docx")

    def test_filename_contains_participants(self, tmp_path):
        path = export_participants_docx(
            [_user(1, "Alice")],
            "TestChat", str(tmp_path),
        )
        assert "_participants_" in os.path.basename(path)

    def test_filename_contains_chat_title(self, tmp_path):
        path = export_participants_docx(
            [_user(1, "Alice")],
            "MyChat", str(tmp_path),
        )
        assert "MyChat" in os.path.basename(path)

    def test_file_is_valid_docx(self, tmp_path):
        path = export_participants_docx(
            [_user(1, "Alice")],
            "TestChat", str(tmp_path),
        )
        doc = Document(path)
        assert len(doc.paragraphs) > 0


class TestParticipantsContent:
    def test_title_heading_present(self, tmp_path):
        path = export_participants_docx(
            [_user(1, "Alice")],
            "TestChat", str(tmp_path),
        )
        doc = Document(path)
        headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
        assert any("TestChat" in h for h in headings)

    def test_metadata_contains_user_count(self, tmp_path):
        path = export_participants_docx(
            [_user(1, "Alice"), _user(2, "Bob")],
            "TestChat", str(tmp_path),
        )
        doc = Document(path)
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "2" in all_text

    def test_user_names_in_table(self, tmp_path):
        path = export_participants_docx(
            [_user(1, "Alice"), _user(2, "Bob")],
            "TestChat", str(tmp_path),
        )
        doc = Document(path)
        table_text = "\n".join(cell.text for row in doc.tables[0].rows for cell in row.cells)
        assert "Alice" in table_text
        assert "Bob" in table_text

    def test_table_has_four_columns_with_counts(self, tmp_path):
        path = export_participants_docx(
            [_user(1, "Alice", message_count=5)],
            "TestChat", str(tmp_path),
        )
        doc = Document(path)
        assert len(doc.tables[0].columns) == 4

    def test_table_has_three_columns_without_counts(self, tmp_path):
        path = export_participants_docx(
            [_user(1, "Alice", message_count=0)],
            "TestChat", str(tmp_path),
        )
        doc = Document(path)
        assert len(doc.tables[0].columns) == 3


class TestParticipantsSorting:
    def test_sorted_by_message_count_descending(self, tmp_path):
        path = export_participants_docx(
            [_user(1, "Alice", message_count=5), _user(2, "Bob", message_count=50)],
            "TestChat", str(tmp_path),
        )
        doc = Document(path)
        rows = doc.tables[0].rows
        bob_row = rows[1].cells[1].text
        alice_row = rows[2].cells[1].text
        assert "Bob" == bob_row
        assert "Alice" == alice_row

    def test_user_with_zero_count_last(self, tmp_path):
        path = export_participants_docx(
            [_user(1, "Zero", message_count=0), _user(2, "Active", message_count=10)],
            "TestChat", str(tmp_path),
        )
        doc = Document(path)
        rows = doc.tables[0].rows
        last_name = rows[-1].cells[1].text
        assert "Zero" == last_name


class TestParticipantsLinks:
    def test_username_creates_t_me_link(self, tmp_path):
        path = export_participants_docx(
            [_user(1, "Alice", "alice_john", 5)],
            "TestChat", str(tmp_path),
        )
        doc = Document(path)
        urls = [str(r.target_ref) for r in doc.part.rels.values()]
        assert any("https://t.me/alice_john" in u for u in urls)

    def test_no_username_creates_tg_link(self, tmp_path):
        path = export_participants_docx(
            [_user(42, "Alice", "", 5)],
            "TestChat", str(tmp_path),
        )
        doc = Document(path)
        doc_xml = doc.element.xml
        assert "tg://user?id=42" in doc_xml


class TestParticipantsEdgeCases:
    def test_empty_user_list(self, tmp_path):
        path = export_participants_docx(
            [], "TestChat", str(tmp_path),
        )
        doc = Document(path)
        assert len(doc.tables[0].rows) == 1  # header only

    def test_unicode_user_names(self, tmp_path):
        path = export_participants_docx(
            [_user(1, "Привет Мир", "user_ru", 5)],
            "TestChat", str(tmp_path),
        )
        doc = Document(path)
        table_text = "\n".join(cell.text for row in doc.tables[0].rows for cell in row.cells)
        assert "Привет Мир" in table_text

    def test_special_chars_in_chat_title(self, tmp_path):
        path = export_participants_docx(
            [_user(1, "Alice")],
            'Chat<>:"|?*', str(tmp_path),
        )
        assert os.path.exists(path)

    def test_single_user(self, tmp_path):
        path = export_participants_docx(
            [_user(1, "Solo", "solo", 3)],
            "TestChat", str(tmp_path),
        )
        doc = Document(path)
        assert len(doc.tables[0].rows) == 2  # header + 1 data row

    def test_output_dir_created_if_missing(self, tmp_path):
        new_dir = str(tmp_path / "nested" / "dir")
        path = export_participants_docx(
            [_user(1, "Alice")],
            "TestChat", new_dir,
        )
        assert os.path.exists(path)
