"""
tests/test_features/test_xml_magic.py

Тесты: xml_magic — bookmarks, internal/external hyperlinks, write_text_with_links.
Использует реальные Document() и Paragraph для проверки XML-структуры.
"""
import pytest
from docx import Document

from features.export import xml_magic


@pytest.fixture(autouse=True)
def fresh_counter():
    xml_magic.reset_counter()
    yield
    xml_magic.reset_counter()


@pytest.fixture
def doc():
    return Document()


@pytest.fixture
def para(doc):
    return doc.add_paragraph()


# ---------------------------------------------------------------------------
# reset_counter
# ---------------------------------------------------------------------------

class TestResetCounter:
    def test_reset_sets_counter_to_zero(self, para):
        xml_magic.add_bookmark(para, "msg_0")
        xml_magic.add_bookmark(para, "msg_1")
        xml_magic.reset_counter()
        xml_magic.add_bookmark(para, "msg_2")
        bm_id = para._p.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}bookmarkStart")
        assert bm_id[-1].get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id") == "0"

    def test_counter_increments_on_bookmark(self, para):
        xml_magic.add_bookmark(para, "msg_0")
        xml_magic.add_bookmark(para, "msg_1")
        starts = para._p.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}bookmarkStart")
        ids = [s.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id") for s in starts]
        assert ids == ["0", "1"]


# ---------------------------------------------------------------------------
# add_bookmark
# ---------------------------------------------------------------------------

class TestAddBookmark:
    def test_bookmark_xml_elements_present(self, para):
        xml_magic.add_bookmark(para, "msg_42")
        xml = para._p.xml
        assert "bookmarkStart" in xml
        assert "bookmarkEnd" in xml

    def test_bookmark_name_set(self, para):
        xml_magic.add_bookmark(para, "msg_42")
        xml = para._p.xml
        assert 'w:name="msg_42"' in xml

    def test_bookmark_id_matches_start_and_end(self, para):
        xml_magic.add_bookmark(para, "msg_1")
        ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        start = para._p.find(f".//{ns}bookmarkStart")
        end = para._p.find(f".//{ns}bookmarkEnd")
        assert start.get(f"{ns}id") == end.get(f"{ns}id")

    def test_multiple_bookmarks_unique_ids(self, para):
        xml_magic.add_bookmark(para, "msg_1")
        xml_magic.add_bookmark(para, "msg_2")
        ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        starts = para._p.findall(f".//{ns}bookmarkStart")
        ids = [s.get(f"{ns}id") for s in starts]
        assert len(set(ids)) == 2


# ---------------------------------------------------------------------------
# add_internal_hyperlink
# ---------------------------------------------------------------------------

class TestAddInternalHyperlink:
    def test_hyperlink_xml_element_created(self, para):
        xml_magic.add_internal_hyperlink(para, 42, "ссылка")
        xml = para._p.xml
        assert "hyperlink" in xml

    def test_anchor_targets_bookmark(self, para):
        xml_magic.add_internal_hyperlink(para, 42, "текст")
        xml = para._p.xml
        assert 'w:anchor="msg_42"' in xml

    def test_link_text_correct(self, para):
        xml_magic.add_internal_hyperlink(para, 42, "сообщение #42")
        xml = para._p.xml
        assert "сообщение #42" in xml

    def test_default_color_blue(self, para):
        xml_magic.add_internal_hyperlink(para, 42, "link")
        xml = para._p.xml
        assert "0563C1" in xml

    def test_custom_color_applied(self, para):
        xml_magic.add_internal_hyperlink(para, 42, "link", color_hex="FF6BC9")
        xml = para._p.xml
        assert "FF6BC9" in xml

    def test_underline_single(self, para):
        xml_magic.add_internal_hyperlink(para, 42, "link")
        xml = para._p.xml
        assert "single" in xml


# ---------------------------------------------------------------------------
# add_external_hyperlink
# ---------------------------------------------------------------------------

class TestAddExternalHyperlink:
    def test_hyperlink_registered_in_relationships(self, doc, para):
        xml_magic.add_external_hyperlink(para, "https://example.com", "Example")
        rels = doc.part.rels
        assert any(
            "https://example.com" in str(r.target_ref)
            for r in rels.values()
        )

    def test_r_id_set_in_xml(self, doc, para):
        xml_magic.add_external_hyperlink(para, "https://example.com", "Example")
        xml = para._p.xml
        assert "r:id=" in xml

    def test_external_link_text_correct(self, doc, para):
        xml_magic.add_external_hyperlink(para, "https://example.com", "Click here")
        xml = para._p.xml
        assert "Click here" in xml

    def test_url_passed_to_relate_to(self, doc, para):
        xml_magic.add_external_hyperlink(para, "https://t.me/channel", "@channel")
        rels = doc.part.rels
        found = any(
            "https://t.me/channel" in str(r.target_ref)
            for r in rels.values()
        )
        assert found


# ---------------------------------------------------------------------------
# write_text_with_links
# ---------------------------------------------------------------------------

class TestWriteTextWithLinks:
    def test_plain_text_no_links(self, para):
        xml_magic.write_text_with_links(para, "Hello world")
        assert len(para.runs) == 1
        assert para.runs[0].text == "Hello world"

    def test_single_url_becomes_link(self, doc, para):
        xml_magic.write_text_with_links(para, "See https://example.com now")
        xml = para._p.xml
        assert "hyperlink" in xml
        assert "See " in para._p.xml
        assert " now" in para._p.xml

    def test_multiple_urls(self, doc, para):
        xml_magic.write_text_with_links(
            para, "a https://a.com b https://b.com c",
        )
        xml = para._p.xml
        assert xml.count("hyperlink") >= 2

    def test_url_only_text(self, doc, para):
        xml_magic.write_text_with_links(para, "https://example.com")
        xml = para._p.xml
        assert "hyperlink" in xml

    def test_empty_string_noop(self, para):
        xml_magic.write_text_with_links(para, "")
        assert len(para.runs) == 0

    def test_url_at_start_and_end(self, doc, para):
        xml_magic.write_text_with_links(
            para, "https://start.com middle https://end.com",
        )
        xml = para._p.xml
        assert xml.count("hyperlink") >= 2

    def test_url_with_query_params(self, doc, para):
        xml_magic.write_text_with_links(
            para, "https://example.com/path?q=1&r=2",
        )
        xml = para._p.xml
        assert "hyperlink" in xml
